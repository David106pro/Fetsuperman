import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from docx import Document
from docx.shared import Pt
import os
from datetime import datetime

class FolderNameCollector:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("文件名称整理工具")
        self.root.geometry("500x200")
        
        # 创建并设置界面
        self.setup_ui()
        
    def setup_ui(self):
        # 文件夹路径选择框
        self.path_frame = tk.Frame(self.root)
        self.path_frame.pack(pady=20)
        
        self.path_var = tk.StringVar()
        self.path_entry = tk.Entry(self.path_frame, textvariable=self.path_var, width=50)
        self.path_entry.pack(side=tk.LEFT, padx=5)
        
        self.browse_btn = tk.Button(self.path_frame, text="选择文件夹", command=self.browse_folder)
        self.browse_btn.pack(side=tk.LEFT)
        
        # 进度条
        self.progress = ttk.Progressbar(self.root, length=400, mode='determinate')
        self.progress.pack(pady=20)
        
        # 开始按钮
        self.start_btn = tk.Button(self.root, text="开始整理", command=self.start_process)
        self.start_btn.pack()
        
    def browse_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.path_var.set(folder_path)
            
    def collect_folder_names(self, path, level=0):
        items = []
        file_count = 0
        try:
            for item in os.listdir(path):
                file_count += 1
                full_path = os.path.join(path, item)
                if os.path.isdir(full_path):
                    items.append({"name": item, "level": level, "path": full_path, "is_dir": True})
                    sub_items, sub_count = self.collect_folder_names(full_path, level + 1)
                    items.extend(sub_items)
                    file_count += sub_count
                else:
                    items.append({"name": item, "level": level, "path": full_path, "is_dir": False})
        except Exception as e:
            print(f"Error accessing {path}: {str(e)}")
        return items, file_count
    
    def write_to_doc(self, doc, items, start_level=0):
        current_level = start_level
        current_line = []
        
        for item in items:
            level = item["level"]
            name = item["name"]
            
            if level == current_level and current_line:
                current_line.append(name)
            else:
                if current_line:
                    p = doc.add_paragraph("; ".join(current_line))
                    # 根据层级设置字体大小
                    font_size = max(16 - (current_level * 2), 8)  # 最小字号为8
                    for run in p.runs:
                        run.font.size = Pt(font_size)
                current_line = [name]
                current_level = level
        
        # 添加最后一行
        if current_line:
            p = doc.add_paragraph("; ".join(current_line))
            font_size = max(16 - (current_level * 2), 8)
            for run in p.runs:
                run.font.size = Pt(font_size)

    def start_process(self):
        try:
            folder_path = self.path_var.get()
            if not folder_path:
                tk.messagebox.showerror("错误", "请先选择文件夹！")
                return

            # 创建主文档
            try:
                main_doc = Document()
                appendix_doc = Document()
            except Exception as e:
                tk.messagebox.showerror("错误", f"创建Word文档失败: {str(e)}\n请确保已安装python-docx库")
                return
            
            # 获取所有文件夹名称和计数
            all_items, _ = self.collect_folder_names(folder_path)
            if not all_items:
                tk.messagebox.showwarning("警告", "所选文件夹为空或无法访问")
                return
                
            total_items = len(all_items)
            
            # 更新进度条
            self.progress['value'] = 0
            self.root.update()

            # 处理大文件夹
            large_folders = {}
            current_folder = None
            folder_items = []
            
            for item in all_items:
                if item["is_dir"]:
                    if current_folder:
                        if len(folder_items) >= 100:
                            large_folders[current_folder] = folder_items
                        folder_items = []
                    current_folder = item["path"]
                folder_items.append(item)
            
            # 检查最后一个文件夹
            if current_folder and len(folder_items) >= 100:
                large_folders[current_folder] = folder_items

            # 写入主文档
            remaining_items = [item for item in all_items if not any(item in folder_items for folder_items in large_folders.values())]
            self.write_to_doc(main_doc, remaining_items)

            # 保存文档
            try:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                main_doc_path = os.path.join(folder_path, f"文件名称整理_{timestamp}.docx")
                main_doc.save(main_doc_path)
                
                if large_folders:
                    appendix_doc.add_heading('附录', level=1)
                    for folder_path, folder_items in large_folders.items():
                        folder_name = os.path.basename(folder_path)
                        appendix_doc.add_heading(folder_name, level=2)
                        self.write_to_doc(appendix_doc, folder_items)
                    
                    appendix_doc_path = os.path.join(folder_path, f"附录_{timestamp}.docx")
                    appendix_doc.save(appendix_doc_path)
                
                tk.messagebox.showinfo("完成", f"文件名称整理完成！\n主文档保存在：{main_doc_path}")
            except Exception as e:
                tk.messagebox.showerror("错误", f"保存文档时出错: {str(e)}")
                return

        except Exception as e:
            tk.messagebox.showerror("错误", f"处理过程中出错: {str(e)}")
        finally:
            self.progress['value'] = 0

    def run(self):
        self.root.mainloop()

if __name__ == "__main__":
    app = FolderNameCollector()
    app.run()
